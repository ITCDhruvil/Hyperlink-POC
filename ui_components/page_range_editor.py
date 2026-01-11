"""
Visual Page Range Editor for Streamlit
Combines auto-extraction and manual table editing
"""
import streamlit as st
from typing import List, Dict, Optional
import re


def parse_range_string(range_str: str) -> Dict:
    """
    Parse a range string like "1-2" or "5" or "25-29, 31-35"

    Returns:
        {
            'start': int,
            'end': Optional[int],
            'is_range': bool,
            'display': str
        }
    """
    range_str = range_str.strip()

    # Check if it contains comma (complex range like "25-29, 31-35")
    if ',' in range_str:
        return {
            'start': None,
            'end': None,
            'is_range': True,
            'is_complex': True,
            'display': range_str
        }

    # Check if it's a range
    if '-' in range_str:
        parts = range_str.split('-')
        if len(parts) == 2 and parts[0].isdigit() and parts[1].isdigit():
            return {
                'start': int(parts[0]),
                'end': int(parts[1]),
                'is_range': True,
                'is_complex': False,
                'display': f"{parts[0]}-{parts[1]}"
            }

    # Single page
    if range_str.isdigit():
        return {
            'start': int(range_str),
            'end': None,
            'is_range': False,
            'is_complex': False,
            'display': range_str
        }

    # Invalid format
    return None


def validate_page_ranges(ranges: List[Dict]) -> List[str]:
    """
    Validate page ranges and return list of errors
    """
    errors = []

    for i, r in enumerate(ranges):
        if not r:
            errors.append(f"Range {i+1}: Invalid format")
            continue

        if r.get('is_complex'):
            continue  # Skip validation for complex ranges

        start = r.get('start')
        end = r.get('end')

        if start is None:
            errors.append(f"Range {i+1}: Missing start page")
            continue

        if start <= 0:
            errors.append(f"Range {i+1}: Page numbers must be positive")

        if end is not None and end < start:
            errors.append(f"Range {i+1}: End page must be >= start page")

    return errors


def format_ranges_for_split(ranges: List[str]) -> str:
    """
    Format list of ranges for PDF split input

    Args:
        ranges: ['1-2', '4', '7-8', '25-29, 31-35']

    Returns:
        '1-2;4;7-8;25-29, 31-35'
    """
    return ';'.join(ranges)


def render_page_range_editor(
    initial_ranges: Optional[List[str]] = None,
    key_prefix: str = "page_range"
) -> List[str]:
    """
    Render interactive page range table editor

    Args:
        initial_ranges: Initial list of ranges to populate table
        key_prefix: Unique prefix for widget keys

    Returns:
        List of page ranges as strings
    """
    # Initialize session state
    if f'{key_prefix}_ranges' not in st.session_state:
        st.session_state[f'{key_prefix}_ranges'] = initial_ranges or []

    ranges = st.session_state[f'{key_prefix}_ranges']

    st.markdown("### 📝 Page Ranges")

    # Display table
    if len(ranges) == 0:
        st.info("No page ranges added yet. Click '+ Add Range' to start.")
    else:
        # Create columns for table header
        cols = st.columns([0.5, 3, 3, 1])
        cols[0].markdown("**#**")
        cols[1].markdown("**Range**")
        cols[2].markdown("**Preview**")
        cols[3].markdown("**Action**")

        st.divider()

        # Render each range row
        ranges_to_remove = []
        updated_ranges = []

        for i, range_str in enumerate(ranges):
            cols = st.columns([0.5, 3, 3, 1])

            # Row number
            cols[0].markdown(f"{i+1}.")

            # Editable range input
            new_range = cols[1].text_input(
                f"Range {i+1}",
                value=range_str,
                key=f"{key_prefix}_range_input_{i}",
                label_visibility="collapsed",
                placeholder="e.g., 1-2 or 5 or 25-29, 31-35"
            )

            # Preview/validation
            parsed = parse_range_string(new_range)
            if parsed:
                if parsed.get('is_complex'):
                    cols[2].markdown(f"✓ Complex: `{parsed['display']}`")
                elif parsed['is_range']:
                    cols[2].markdown(f"✓ Pages {parsed['start']} to {parsed['end']}")
                else:
                    cols[2].markdown(f"✓ Single page {parsed['start']}")
            else:
                cols[2].markdown("❌ Invalid format")

            # Remove button
            if cols[3].button("🗑️", key=f"{key_prefix}_remove_{i}"):
                ranges_to_remove.append(i)

            updated_ranges.append(new_range)

        # Remove marked ranges
        if ranges_to_remove:
            updated_ranges = [r for i, r in enumerate(updated_ranges) if i not in ranges_to_remove]
            st.session_state[f'{key_prefix}_ranges'] = updated_ranges
            st.rerun()
        else:
            st.session_state[f'{key_prefix}_ranges'] = updated_ranges

    # Add new range button
    col1, col2, col3 = st.columns([2, 2, 6])
    if col1.button("➕ Add Range", key=f"{key_prefix}_add"):
        st.session_state[f'{key_prefix}_ranges'].append("")
        st.rerun()

    if col2.button("🗑️ Clear All", key=f"{key_prefix}_clear"):
        st.session_state[f'{key_prefix}_ranges'] = []
        st.rerun()

    # Validation summary
    current_ranges = st.session_state[f'{key_prefix}_ranges']
    if current_ranges:
        parsed_ranges = [parse_range_string(r) for r in current_ranges if r.strip()]
        errors = validate_page_ranges(parsed_ranges)

        if errors:
            st.error("⚠️ Validation Errors:")
            for error in errors:
                st.markdown(f"- {error}")
        else:
            valid_ranges = [r for r in current_ranges if r.strip()]
            if valid_ranges:
                formatted = format_ranges_for_split(valid_ranges)
                st.success(f"✓ {len(valid_ranges)} ranges ready")

                with st.expander("📋 Formatted for Split"):
                    st.code(formatted, language=None)

    return [r for r in current_ranges if r.strip()]


def auto_extract_from_word(docx_file) -> List[str]:
    """
    Auto-extract page ranges from uploaded Word document

    Args:
        docx_file: Streamlit UploadedFile object

    Returns:
        List of page ranges ['1-2', '4', '7-8', ...]
    """
    from docx import Document
    import tempfile
    import os

    # Save uploaded file temporarily
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
        tmp_file.write(docx_file.getvalue())
        tmp_path = tmp_file.name

    try:
        # Embedded extraction logic to avoid import dependencies
        def _is_statement_line(text: str) -> bool:
            text = (text or '').replace('\u00a0', ' ').strip()
            return bool(re.match(r'^\d{1,2}/\d{1,2}/\d{2,4}\.\s+', text))

        def _normalize_page_spec(s: str) -> str:
            s = (s or '').replace('\u00a0', ' ').strip()
            s = s.replace('–', '-').replace('—', '-').replace('‑', '-')
            s = re.sub(r'(\d+)\s*-\s*(\d+)', r'\1-\2', s)
            s = re.sub(r'\s+', ' ', s)
            s = re.sub(r'\s*,\s*', ', ', s)
            return s.strip()

        def _is_page_spec(s: str) -> bool:
            return bool(re.fullmatch(r'\d+(?:\s*-\s*\d+)?(?:\s*,\s*\d+(?:\s*-\s*\d+)?)*', (s or '').strip()))

        def _parse_statement_with_page_number(text: str):
            text = _normalize_page_spec(text)
            page_spec_pattern = re.compile(r'(?P<pages>\d+(?:\s*-\s*\d+)?(?:\s*,\s*\d+(?:\s*-\s*\d+)?)*)')

            best = None
            for m in page_spec_pattern.finditer(text):
                start, end = m.start('pages'), m.end('pages')

                if end < len(text) and text[end] == '/':
                    continue
                if start > 0 and text[start - 1] == '/':
                    continue
                if start > 0 and text[start - 1] == '.':
                    continue
                if end < len(text) and text[end] == ')':
                    continue

                pages_raw = (m.group('pages') or '').strip()
                if not _is_page_spec(pages_raw):
                    continue

                if start > 0:
                    before = text[:start].rstrip()
                    if not before:
                        continue
                    if not (before.endswith('.') or before.endswith(':') or before.endswith(' ')):
                        continue

                    context_before = before[-15:] if len(before) >= 15 else before
                    if re.search(r'(SpO2|BP|Pulse|Temp|Weight|RR|P|T):\s*$', context_before, re.IGNORECASE):
                        continue

                if start > 250:
                    if best is not None:
                        continue

                if best is None:
                    best = m

            if best:
                page_spec_raw = (best.group('pages') or '').strip()
                page_spec = _normalize_page_spec(page_spec_raw)
                return {'page_range': page_spec}
            return None

        doc = Document(tmp_path)

        all_paragraphs = []
        for p in doc.paragraphs:
            all_paragraphs.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        all_paragraphs.append(p.text)

        page_ranges = []
        seen_ranges = set()

        for text in all_paragraphs:
            if not text or not text.strip():
                continue

            if not _is_statement_line(text):
                continue

            result = _parse_statement_with_page_number(text)
            if not result or not result.get('page_range'):
                continue

            page_range = result['page_range']

            if page_range not in seen_ranges:
                page_ranges.append(page_range)
                seen_ranges.add(page_range)

        return page_ranges

    finally:
        # Clean up temp file
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)


if __name__ == '__main__':
    # Example usage
    st.set_page_config(page_title="Page Range Editor Demo", layout="wide")

    st.title("Page Range Editor Demo")

    # Demo with some initial ranges
    ranges = render_page_range_editor(
        initial_ranges=['1-2', '4', '7-8', '11-15'],
        key_prefix="demo"
    )

    st.markdown("---")
    st.markdown("### Current Ranges")
    st.write(ranges)
