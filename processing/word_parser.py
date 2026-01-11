"""
Parse Word documents to extract medical record statements and attach Drive links
"""
import re
from pathlib import Path
from typing import List, Dict, Tuple, Optional
from docx import Document
from datetime import datetime


def parse_medical_statement(text: str) -> Optional[Dict]:
    """
    Parse a medical record statement to extract components

    Examples:
    - "03/05/21. Emergency Department Record. KAISER PERMANENTE. S, page No. 9-23"
    - "04/05/13.  Progress Note.  US HEALTHWORKS MEDICAL GROUP"
    - "10/17/2025.  Attestation.  From CONCENTRA-ONTARI."

    Returns:
        Dict with: date, doc_type, facility, page_range (if found)
    """
    # Pattern: DATE.  DOCUMENT_TYPE.  FACILITY.  [optional text], [optional page range]
    # Date formats: MM/DD/YY or MM/DD/YYYY

    # Try to extract date (at the beginning)
    date_match = re.match(r'(\d{1,2}/\d{1,2}/\d{2,4})', text.strip())
    if not date_match:
        return None

    date_str = date_match.group(1)
    remaining_text = text[len(date_str):].strip()

    # Remove leading period/dot
    if remaining_text.startswith('.'):
        remaining_text = remaining_text[1:].strip()

    # Split by periods to get segments
    segments = [s.strip() for s in remaining_text.split('.') if s.strip()]

    if len(segments) < 2:
        return None

    # First segment is usually the document type
    doc_type = segments[0]

    # Second segment is usually the facility (may have "From" prefix)
    facility = segments[1]
    if facility.lower().startswith('from '):
        facility = facility[5:].strip()

    # Look for page range in the entire text
    # Patterns: "page No. 9-23", "9-23", "pages 9-23", etc.
    page_range = None
    page_match = re.search(r'(?:page(?:s)?\s*(?:No\.?|#)?\s*)?(\d+)\s*-\s*(\d+)', text, re.IGNORECASE)
    if page_match:
        start_page = page_match.group(1)
        end_page = page_match.group(2)
        page_range = f"{start_page}-{end_page}"
    else:
        # Try single page
        single_page_match = re.search(r'(?:page\s*(?:No\.?|#)?\s*)?(\d+)(?!\s*-)', text, re.IGNORECASE)
        if single_page_match:
            page_num = single_page_match.group(1)
            page_range = f"{page_num}-{page_num}"

    return {
        'date': date_str,
        'doc_type': doc_type,
        'facility': facility,
        'page_range': page_range,
        'original_text': text.strip()
    }


def extract_patient_name_from_docx(doc: Document) -> Optional[str]:
    """
    Extract patient name from Word document
    Looks for patterns like "PATIENT NAME: Carl Mayfield"
    """
    for paragraph in doc.paragraphs[:20]:  # Check first 20 paragraphs
        text = paragraph.text.strip()

        # Pattern: "PATIENT NAME:" followed by the name
        match = re.search(r'PATIENT\s*NAME[:\s]+([A-Za-z\s\.]+?)(?:\s*$|\s*DATE)', text, re.IGNORECASE)
        if match:
            return match.group(1).strip()

    return None


def find_matching_pdf_in_database(patient_name: str, page_range: str, date_str: str) -> Optional[Tuple[str, str]]:
    """
    Find matching PDF in database using patient name, page range, and date

    Returns:
        Tuple of (drive_link, local_path) if found, None otherwise
    """
    from pdfs.models import PDFSet, Patient
    from datetime import datetime

    # Parse date to help narrow search
    try:
        # Try different date formats
        for date_format in ['%m/%d/%Y', '%m/%d/%y']:
            try:
                parsed_date = datetime.strptime(date_str, date_format).date()
                break
            except:
                continue
    except:
        parsed_date = None

    # Parse page range
    try:
        if '-' in page_range:
            start_page, end_page = map(int, page_range.split('-'))
        else:
            start_page = end_page = int(page_range)
    except:
        return None

    # Find patient by name (fuzzy match)
    patients = Patient.objects.filter(name__icontains=patient_name)

    if not patients.exists():
        # Try without middle name or with name variations
        name_parts = patient_name.split()
        if len(name_parts) >= 2:
            # Try first and last name only
            first_last = f"{name_parts[0]} {name_parts[-1]}"
            patients = Patient.objects.filter(name__icontains=first_last)

    if not patients.exists():
        return None

    # Search for PDF sets matching the page range
    for patient in patients:
        pdf_sets = PDFSet.objects.filter(
            patient=patient,
            start_page=start_page,
            end_page=end_page,
            state='UPLOADED'
        )

        # If date is available, filter by date
        if parsed_date:
            pdf_sets = pdf_sets.filter(date=parsed_date)

        if pdf_sets.exists():
            pdf_set = pdf_sets.first()
            if pdf_set.drive_webview_link:
                return (pdf_set.drive_webview_link, pdf_set.local_path)

    # If exact match not found, try approximate match (within +/-1 page)
    for patient in patients:
        pdf_sets = PDFSet.objects.filter(
            patient=patient,
            start_page__gte=start_page - 1,
            start_page__lte=start_page + 1,
            end_page__gte=end_page - 1,
            end_page__lte=end_page + 1,
            state='UPLOADED'
        )

        if pdf_sets.exists():
            pdf_set = pdf_sets.first()
            if pdf_set.drive_webview_link:
                return (pdf_set.drive_webview_link, pdf_set.local_path)

    return None


def process_word_document_with_links(
    input_docx_path: str,
    output_docx_path: str,
    patient_name: Optional[str] = None
) -> Dict:
    """
    Process a Word document to find medical statements and attach Drive links

    Args:
        input_docx_path: Path to input Word document
        output_docx_path: Path to save output Word document
        patient_name: Optional patient name (if not auto-detected)

    Returns:
        Dict with statistics: {
            'total_statements': int,
            'linked_statements': int,
            'unlinked_statements': int,
            'patient_name': str
        }
    """
    from processing.docx_utils import add_hyperlink

    doc = Document(input_docx_path)

    # Auto-detect patient name if not provided
    if not patient_name:
        patient_name = extract_patient_name_from_docx(doc)

    if not patient_name:
        raise ValueError("Patient name not found in document and not provided")

    stats = {
        'total_statements': 0,
        'linked_statements': 0,
        'unlinked_statements': 0,
        'patient_name': patient_name,
        'statements': []
    }

    # Process each paragraph
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        # Skip empty paragraphs
        if not text:
            continue

        # Try to parse as medical statement
        statement = parse_medical_statement(text)

        if statement and statement.get('page_range'):
            stats['total_statements'] += 1

            # Find matching PDF
            match = find_matching_pdf_in_database(
                patient_name,
                statement['page_range'],
                statement['date']
            )

            if match:
                drive_link, local_path = match
                stats['linked_statements'] += 1

                # Clear existing text and add hyperlink
                paragraph.clear()
                add_hyperlink(
                    paragraph,
                    drive_link,
                    statement['original_text'],
                    color='0563C1',
                    underline=True
                )

                # Make it bold
                for run in paragraph.runs:
                    run.bold = True

                stats['statements'].append({
                    'text': statement['original_text'],
                    'page_range': statement['page_range'],
                    'status': 'linked',
                    'drive_link': drive_link
                })
            else:
                stats['unlinked_statements'] += 1
                stats['statements'].append({
                    'text': statement['original_text'],
                    'page_range': statement['page_range'],
                    'status': 'not_found',
                    'reason': 'No matching PDF found in database'
                })

    # Save the modified document
    doc.save(output_docx_path)

    return stats


def test_statement_parser():
    """Test the statement parser with example statements"""
    test_statements = [
        "10/17/2025.  Attestation.  From CONCENTRA-ONTARI.",
        "10/17/25.  Attestation.   From KAISER PERMANENTE MEDICAL CENTER.",
        "04/05/13.  Progress Note.  US HEALTHWORKS MEDICAL GROUP",
        "04/10/13.  Primary Treating Physician's Progress Report (PR-2).  US HEALTHWORKS MEDICAL GROUP",
        "03/05/21. Emergency Department Record. KAISER PERMANENTE. S, page No. 9-23",
        "04/14/21. Doctor's First Report of Occupational Injury or Illness. KAISER PERMANENTE. 17-26",
        "Mohammad Mahmud, MD  20-23 OT_8896048_ME_Records_001",
    ]

    print("Testing Medical Statement Parser:")
    print("=" * 80)

    for stmt in test_statements:
        result = parse_medical_statement(stmt)
        print(f"\nInput: {stmt}")
        if result:
            print(f"  Date: {result['date']}")
            print(f"  Type: {result['doc_type']}")
            print(f"  Facility: {result['facility']}")
            print(f"  Pages: {result.get('page_range', 'Not found')}")
        else:
            print("  [FAILED TO PARSE]")

    print("\n" + "=" * 80)


if __name__ == "__main__":
    test_statement_parser()
