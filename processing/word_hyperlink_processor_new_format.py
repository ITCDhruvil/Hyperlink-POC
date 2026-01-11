"""
New Format Word Hyperlink Processor

NEW BEHAVIOR:
- Extract page range from doctor statement line
- Insert hyperlink in the DESCRIPTION line (next line)
- Keep doctor line but remove page range
"""
import re
from typing import Optional, Dict, List
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .drive_utils import get_drive_service


def add_hyperlink(paragraph, url, text, color='0000FF', underline=True, bold=False):
    """Add a hyperlink to a paragraph using OXML"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)

    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    new_run.append(rPr)

    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


class WordHyperlinkProcessorNewFormat:
    """
    NEW FORMAT: Hyperlink description lines, not doctor lines
    """

    def __init__(self):
        self.drive_service = get_drive_service()

    def extract_patient_name_from_document(self, doc: Document) -> Optional[str]:
        """
        Extract patient name from top of document

        Looks for patterns like:
        - "PATIENT NAME: CARL MAYFIELD"
        - "Patient: Carl Mayfield"
        - "Name: Carl Mayfield"
        """
        # Check first 10 paragraphs
        for paragraph in doc.paragraphs[:10]:
            text = paragraph.text.strip()

            # Pattern 1: "PATIENT NAME: CARL MAYFIELD"
            match = re.search(r'PATIENT\s+NAME:?\s+(.+)', text, re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                # Convert to folder format: "Carl Mayfield" -> "Carl_Mayfield"
                name = name.title().replace(' ', '_')
                return name

            # Pattern 2: "Name: Carl Mayfield"
            match = re.search(r'^Name:?\s+(.+)', text, re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                name = name.title().replace(' ', '_')
                return name

        return None

    def parse_doctor_statement(self, text: str) -> Optional[Dict]:
        """
        Parse doctor statement to extract page range

        Format: "Doctor Name  PageRange  FileIdentifier"
        Example: "Mohammad Mahmud, MD  20-23 OT_8896048_ME_Records_001"

        Returns page range for linking the next line
        """
        pattern = r'^(.*?)\s+(\d+\s*-\s*\d+(?:\s*,\s*\d+\s*-\s*\d+)*)\s*-?\s*(OT_\S+)\s*$'
        match = re.match(pattern, text.strip())

        if match:
            doctor_name = match.group(1).strip()
            page_range_str = match.group(2).strip()
            file_identifier = match.group(3).strip()

            # Parse page ranges
            page_ranges = [r.strip() for r in page_range_str.split(',')]
            primary_page_range = page_ranges[0]

            # Clean text for doctor line (without page range)
            clean_text = f"{doctor_name} - {file_identifier}"

            return {
                'page_range': primary_page_range,
                'all_page_ranges': page_ranges,
                'clean_text': clean_text,
                'original_text': text.strip(),
                'doctor_name': doctor_name,
                'file_id': file_identifier
            }

        return None

    def is_description_line(self, text: str) -> bool:
        """
        Check if line is a description line

        Patterns:
        - Starts with date: "04/05/13. Progress Note..."
        - Contains keywords: "Progress Note", "Report", "Record"
        """
        # Pattern: starts with date MM/DD/YY or MM/DD/YYYY
        if re.match(r'^\d{1,2}/\d{1,2}/\d{2,4}\.', text):
            return True

        return False

    def get_pdfs_from_drive_folder(self, drive_folder_id: str) -> Dict[str, str]:
        """Get all PDF files from Drive folder"""
        pdf_links = {}

        try:
            results = self.drive_service.service.files().list(
                q=f"'{drive_folder_id}' in parents and mimeType='application/pdf' and trashed=false",
                fields="files(id, name, webViewLink)",
                pageSize=1000,
                includeItemsFromAllDrives=True,
                supportsAllDrives=True
            ).execute()

            files = results.get('files', [])

            for file in files:
                filename = file['name']
                if filename.endswith('.pdf'):
                    page_range = filename[:-4]
                    if re.match(r'^\d+-\d+$', page_range):
                        pdf_links[page_range] = file['webViewLink']

            return pdf_links

        except Exception as e:
            raise Exception(f"Failed to fetch PDFs: {str(e)}")

    def process_word_document(
        self,
        input_docx_path: str,
        pdf_links: Dict[str, str],
        output_docx_path: Optional[str] = None
    ) -> Dict:
        """
        Process Word document - NEW FORMAT

        Logic:
        1. Find doctor statement line -> extract page range
        2. Next line (description) -> insert hyperlink
        3. Clean up doctor line (remove page range)
        """
        if output_docx_path is None:
            output_docx_path = input_docx_path

        doc = Document(input_docx_path)

        stats = {
            'total_statements': 0,
            'linked_statements': 0,
            'unlinked_statements': 0,
            'statements': []
        }

        paragraphs = doc.paragraphs
        i = 0

        while i < len(paragraphs):
            paragraph = paragraphs[i]
            text = paragraph.text.strip()

            if not text:
                i += 1
                continue

            # Try to parse as doctor statement
            doctor_info = self.parse_doctor_statement(text)

            if doctor_info:
                stats['total_statements'] += 1
                page_range = doctor_info['page_range']
                clean_text = doctor_info['clean_text']

                # Update current paragraph (doctor line) - remove page range
                paragraph.clear()
                run = paragraph.add_run(clean_text)
                run.bold = True

                # Check if next line is description
                if i + 1 < len(paragraphs):
                    next_paragraph = paragraphs[i + 1]
                    next_text = next_paragraph.text.strip()

                    if next_text and self.is_description_line(next_text):
                        # Insert hyperlink in description line
                        if page_range in pdf_links:
                            drive_link = pdf_links[page_range]
                            stats['linked_statements'] += 1

                            # Clear and add hyperlink
                            next_paragraph.clear()
                            add_hyperlink(
                                next_paragraph,
                                drive_link,
                                next_text,
                                color='0563C1',
                                underline=True,
                                bold=False
                            )

                            stats['statements'].append({
                                'page_range': page_range,
                                'doctor_line': clean_text,
                                'description_line': next_text,
                                'status': 'linked',
                                'drive_link': drive_link
                            })
                        else:
                            stats['unlinked_statements'] += 1
                            stats['statements'].append({
                                'page_range': page_range,
                                'doctor_line': clean_text,
                                'description_line': next_text,
                                'status': 'not_found',
                                'reason': f'No PDF found for pages {page_range}'
                            })

            i += 1

        # Save document
        doc.save(output_docx_path)

        return stats
