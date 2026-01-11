"""
Extract page ranges from Word documents for PDF splitting
"""
from docx import Document
from typing import List
import re


def extract_page_ranges_from_word(docx_path: str) -> List[str]:
    """
    Extract all page ranges from a Word document.

    Looks for statements starting with dates and extracts the page numbers.

    Args:
        docx_path: Path to Word document

    Returns:
        List of page ranges in order (e.g., ['1-2', '4', '7-8', '11-15', ...])

    Example:
        Input document with:
            06/19/25. From HOSPITAL. Report. 1-2
            07/30/25. From CLINIC. Report. 4
            08/26/22. From LAB. Report. 7-8

        Returns:
            ['1-2', '4', '7-8']
    """
    from .word_hyperlink_processor_simple import WordHyperlinkProcessorSimple

    processor = WordHyperlinkProcessorSimple()
    doc = Document(docx_path)

    page_ranges = []
    seen_ranges = set()  # Track duplicates

    # Get all paragraphs (body + tables)
    all_paragraphs = []
    for p in doc.paragraphs:
        all_paragraphs.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    all_paragraphs.append(p.text)

    # Process each paragraph
    for text in all_paragraphs:
        if not text or not text.strip():
            continue

        # Check if it's a statement line
        if not processor.is_statement_line(text):
            continue

        # Extract page range
        result = processor.parse_statement_with_page_number(text)
        if not result or not result.get('page_range'):
            continue

        page_range = result['page_range']

        # Avoid duplicates (in case same range appears multiple times)
        if page_range not in seen_ranges:
            page_ranges.append(page_range)
            seen_ranges.add(page_range)

    return page_ranges


def format_for_split_input(page_ranges: List[str]) -> str:
    """
    Format page ranges for PDF split input.

    Args:
        page_ranges: List of page ranges ['1-2', '4', '7-8']

    Returns:
        Semicolon-separated string: '1-2;4;7-8'
    """
    return ';'.join(page_ranges)


def preview_page_ranges(docx_path: str) -> dict:
    """
    Extract and preview page ranges from Word document.

    Returns:
        {
            'page_ranges': ['1-2', '4', '7-8', ...],
            'count': 13,
            'formatted': '1-2;4;7-8;...',
            'preview': [
                {'range': '1-2', 'description': '06/19/25. From HOSPITAL...'},
                {'range': '4', 'description': '07/30/25. From CLINIC...'},
                ...
            ]
        }
    """
    from .word_hyperlink_processor_simple import WordHyperlinkProcessorSimple

    processor = WordHyperlinkProcessorSimple()
    doc = Document(docx_path)

    page_ranges = []
    preview_items = []
    seen_ranges = set()

    # Get all paragraphs
    all_paragraphs = []
    for p in doc.paragraphs:
        all_paragraphs.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    all_paragraphs.append(p.text)

    # Process each paragraph
    for text in all_paragraphs:
        if not text or not text.strip():
            continue

        if not processor.is_statement_line(text):
            continue

        result = processor.parse_statement_with_page_number(text)
        if not result or not result.get('page_range'):
            continue

        page_range = result['page_range']

        if page_range not in seen_ranges:
            page_ranges.append(page_range)
            seen_ranges.add(page_range)

            # Create preview with truncated description
            description = result.get('header_text', '')
            if len(description) > 60:
                description = description[:60] + '...'

            preview_items.append({
                'range': page_range,
                'description': description
            })

    return {
        'page_ranges': page_ranges,
        'count': len(page_ranges),
        'formatted': format_for_split_input(page_ranges),
        'preview': preview_items
    }


if __name__ == '__main__':
    # Test with sample document
    import sys

    if len(sys.argv) > 1:
        docx_path = sys.argv[1]
    else:
        docx_path = r"D:\hyperlink_POC\sample 2\Al Yabroudi, Ahmad_ROR_Dr. Devidson_11-24-25----1 Point 1 output (1).docx"

    print("=" * 80)
    print("PAGE RANGE EXTRACTOR - PREVIEW")
    print("=" * 80)

    result = preview_page_ranges(docx_path)

    print(f"\nDocument: {docx_path}")
    print(f"Total Page Ranges Found: {result['count']}")
    print()

    print("Preview:")
    print("-" * 80)
    for item in result['preview']:
        print(f"  [{item['range']:20s}] {item['description']}")

    print()
    print("Formatted for Split Input:")
    print("-" * 80)
    print(result['formatted'])
    print()

    print("=" * 80)
    print("Copy the formatted string above and paste it into your PDF split tool!")
    print("=" * 80)
