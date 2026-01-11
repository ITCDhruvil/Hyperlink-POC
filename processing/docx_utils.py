"""
Word document generation with hyperlinks
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import List, Dict
from pathlib import Path
from .drive_utils import get_drive_service


def add_hyperlink(paragraph, url, text, color='0000FF', underline=True, bold=False):
    """
    Add a hyperlink to a paragraph using OXML
    Args:
        paragraph: docx paragraph object
        url: URL to link to
        text: Display text for the link
        color: Hex color code (default blue)
        underline: Whether to underline the link
        bold: Whether to make the text bold
    """
    # Create relationship to external URL
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    # Create hyperlink element
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    # Create run element
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Apply Word's built-in "Hyperlink" character style
    # This ensures proper blue underline rendering in Word
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    # Add color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)

    # Add underline if needed
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    # Add bold if needed
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    new_run.append(rPr)

    # Add text
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def verify_drive_link(file_id: str) -> bool:
    """
    Verify that a Drive link is accessible
    Returns True if accessible, False otherwise
    """
    try:
        drive_service = get_drive_service()
        return drive_service.verify_file_exists(file_id)
    except Exception as e:
        print(f"Error verifying Drive link: {e}")
        return False


def generate_summary_document(patient_data: Dict, pdf_sets: List[Dict], output_path: str) -> str:
    """
    Generate a Word document with patient summary and hyperlinks
    
    Args:
        patient_data: Dict with patient info (name, patient_id, address, contact)
        pdf_sets: List of dicts with PDF set info (date, drive_webview_link, drive_file_id)
        output_path: Path to save the document
    
    Returns:
        Path to generated document
    """
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Add title
    title = doc.add_heading('Patient Medical Records Summary', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add patient information section
    doc.add_heading('Patient Information', level=1)
    
    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Light Grid Accent 1'
    
    # Patient ID
    info_table.rows[0].cells[0].text = 'Patient ID:'
    info_table.rows[0].cells[1].text = patient_data.get('patient_id', 'N/A')
    
    # Name
    info_table.rows[1].cells[0].text = 'Name:'
    info_table.rows[1].cells[1].text = patient_data.get('name', 'N/A')
    
    # Address
    info_table.rows[2].cells[0].text = 'Address:'
    info_table.rows[2].cells[1].text = patient_data.get('address', 'N/A')
    
    # Contact
    info_table.rows[3].cells[0].text = 'Contact:'
    info_table.rows[3].cells[1].text = patient_data.get('contact', 'N/A')
    
    # Add spacing
    doc.add_paragraph()
    
    # Add medical records section
    doc.add_heading('MEDICAL RECORDS:', level=1)
    
    # Add each PDF set with hyperlink
    for idx, pdf_set in enumerate(pdf_sets, 1):
        # 1. HEADER LINE (Hyperlinked)
        p_header = doc.add_paragraph()
        
        # Build Header Text: "Doctor Name  PageRange  OriginalFile"
        date_str = pdf_set.get('date', 'N/A')
        hospital = pdf_set.get('hospital', 'Unknown Hospital')
        start_page = pdf_set.get('start_page', '?')
        end_page = pdf_set.get('end_page', '?')
        original_filename = pdf_set.get('original_filename', 'Doc')
        doctor_name = pdf_set.get('doctor_name', '')  # Extract from summary if available

        # Format: "Doctor Name  PageRange  OriginalFile" or just "PageRange  OriginalFile" if no doctor
        if doctor_name:
            header_text = f"{doctor_name}  {start_page}-{end_page}  {original_filename}"
        else:
            header_text = f"{start_page}-{end_page}  {original_filename}"
        
        # Verify link before adding
        drive_file_id = pdf_set.get('drive_file_id', '')
        drive_link = pdf_set.get('drive_webview_link', '')
        
        if drive_link and drive_file_id and verify_drive_link(drive_file_id):
            # Add HYPERLINK to the Header - Blue and underlined
            add_hyperlink(p_header, drive_link, header_text, color='0563C1', underline=True)
        else:
            run = p_header.add_run(header_text)
            run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Make Header Bold
        for run in p_header.runs:
            run.bold = True
            
        # 2. BODY LINE (Summary) - Make it bold
        p_body = doc.add_paragraph()
        
        # Get the actual summary if we have the PDF path
        summary_text = pdf_set.get('summary', '')
        if not summary_text:
            # Fallback placeholder
            summary_text = f"{date_str}.  Medical Record.  {hospital}.  (Summary content would appear here extracted from text...)"
        
        run = p_body.add_run(summary_text)
        run.bold = True
        
        # Add spacing after item
        doc.add_paragraph()
    
    # Add footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Generated automatically by PDF Automation System')
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(128, 128, 128)
    footer_run.italic = True
    
    # Ensure output directory exists
    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    
    # Save document
    doc.save(output_path)
    
    return output_path


def generate_patient_summary(patient, pdf_sets_queryset, output_dir: str) -> str:
    """
    Generate summary document for a patient from Django models
    
    Args:
        patient: Patient model instance
        pdf_sets_queryset: QuerySet of PDFSet instances
        output_dir: Directory to save the document
    
    Returns:
        Path to generated document
    """
    from .pdf_utils import generate_section_summary
    from pathlib import Path as PathLib
    
    # Prepare patient data
    patient_data = {
        'patient_id': patient.patient_id,
        'name': patient.name,
        'address': patient.address,
        'contact': patient.contact
    }
    
    # Prepare PDF sets data
    pdf_sets = []
    for pdf_set in pdf_sets_queryset:
        # Generate summary for this section
        summary = "(Summary unavailable)"
        doctor_name = ""
        if pdf_set.original_pdf and pdf_set.original_pdf.file_path:
            try:
                # Get the full path to the PDF
                pdf_path = PathLib(pdf_set.original_pdf.file_path.path)
                summary, doctor_name = generate_section_summary(
                    str(pdf_path),
                    pdf_set.start_page,
                    pdf_set.end_page,
                    max_length=400
                )
            except Exception as e:
                summary = f"(Error: {str(e)[:50]})"
                doctor_name = ""
        
        pdf_sets.append({
            'date': pdf_set.date.strftime('%m/%d/%Y') if pdf_set.date else 'N/A',
            'drive_webview_link': pdf_set.drive_webview_link,
            'drive_file_id': pdf_set.drive_file_id,
            'hospital': pdf_set.hospital,
            'start_page': pdf_set.start_page,
            'end_page': pdf_set.end_page,
            'original_filename': pdf_set.original_pdf.filename if pdf_set.original_pdf else 'Unknown PDF',
            'summary': summary,
            'doctor_name': doctor_name
        })
    
    # Generate filename with timestamp to avoid file locking issues
    import time
    timestamp = int(time.time())
    filename = f"summary_{patient.patient_id}_{patient.name.replace(' ', '_')}_{timestamp}.docx"
    output_path = Path(output_dir) / filename
    
    # Generate document
    return generate_summary_document(patient_data, pdf_sets, str(output_path))
