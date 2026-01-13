"""
Word Document Hyperlink Processor
Generalized module for attaching Google Drive PDF hyperlinks to Word document statements
"""
import re
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from docx import Document
from .docx_utils import add_hyperlink
from .drive_utils import get_drive_service


class WordHyperlinkProcessor:
    """
    Process Word documents to attach Google Drive PDF hyperlinks to statements
    """

    def __init__(self):
        self.drive_service = get_drive_service()

    def parse_statement_with_page_number(self, text: str) -> Optional[Dict]:
        """
        FLEXIBLE parser that extracts page ranges from ANY format

        Format 1: Doctor statement
        - "Jeffrey Chao, MD (Emergency Medicine) 51-66 - OT_8896048_ME_Records_001"

        Format 2: ANY format with page ranges
        - "10/17/25. Attestation. From KAISER. Page No 3-4"
        - "10/17/25. Attestation. From KAISER. 3-4"
        - "10/17/25. Attestation. From KAISER. (Page No 3-4)"
        - "10/17/25. Attestation. From KAISER. (3-4)"
        - "10/17/25. Attestation. From KAISER. Page-No-3-4"
        - "10/17/25. Attestation. From KAISER. Page_No_3-4"
        - And many more variations!

        Args:
            text: Statement text to parse

        Returns:
            Dict with page_range, clean_text, original_text or None if no match
        """
        # Try Format 1: Doctor Name + Page Range + File Identifier
        doctor_pattern = r'^(.*?)\s+(\d+\s*-\s*\d+(?:\s*,\s*\d+\s*-\s*\d+)*)\s*-?\s*(OT_\S+)\s*$'
        doctor_match = re.match(doctor_pattern, text.strip())

        if doctor_match:
            doctor_name = doctor_match.group(1).strip()
            page_range_str = doctor_match.group(2).strip()
            file_identifier = doctor_match.group(3).strip()

            # Parse individual page ranges: "25-27, 29-31" -> ["25-27", "29-31"]
            page_ranges = [r.strip() for r in page_range_str.split(',')]
            primary_page_range = page_ranges[0]

            # Clean text: doctor name + file identifier (page range removed)
            clean_text = f"{doctor_name} - {file_identifier}"

            return {
                'page_range': primary_page_range,
                'all_page_ranges': page_ranges,
                'clean_text': clean_text,
                'original_text': text.strip(),
                'format': 'doctor_statement'
            }

        # Try Format 2: Find ANY page range pattern (digits-digits)
        # This matches ANY format: "Page No 3-4", "3-4", "(3-4)", "Page_No_3-4", etc.
        page_range_pattern = r'\d+\s*-\s*\d+'
        page_matches = list(re.finditer(page_range_pattern, text))

        if page_matches:
            # Use the LAST match (usually at the end of the statement)
            # This avoids matching dates like "10/17/25"
            last_match = page_matches[-1]

            # Extract and normalize page range
            page_range = last_match.group(0).strip()
            page_range = re.sub(r'\s*-\s*', '-', page_range)  # Remove spaces around dash

            # Find where to cut the text
            match_start = last_match.start()
            cutoff_point = match_start

            # Look backwards for page-related keywords or brackets
            text_before_match = text[:match_start]
            search_start = max(0, match_start - 20)
            search_text = text[search_start:match_start].lower()

            # Search for page keywords
            page_keyword_pos = -1
            for keyword in ['page', 'pg', 'p.', 'pages']:
                pos = search_text.rfind(keyword)
                if pos != -1:
                    page_keyword_pos = search_start + pos
                    break

            # Search for opening brackets
            bracket_pos = -1
            for bracket in ['(', '[', '{']:
                pos = text_before_match.rfind(bracket)
                if pos != -1 and pos > page_keyword_pos:
                    bracket_pos = pos

            # Use the earliest marker found
            if page_keyword_pos != -1:
                cutoff_point = page_keyword_pos
            elif bracket_pos != -1:
                cutoff_point = bracket_pos

            # Clean text: remove from cutoff point to end
            clean_text = text[:cutoff_point].strip()
            clean_text = clean_text.rstrip('.,;:([{').strip()

            return {
                'page_range': page_range,
                'clean_text': clean_text,
                'original_text': text.strip(),
                'format': 'flexible_page_range'
            }

        return None

    def get_pdfs_from_drive_folder(self, drive_folder_id: str) -> Dict[str, str]:
        """
        Get all PDF files from a Google Drive folder and map page ranges to webview links

        Args:
            drive_folder_id: Google Drive folder ID containing PDFs

        Returns:
            Dict mapping page_range -> drive_webview_link
            Example: {'3-4': 'https://drive.google.com/file/d/...', ...}
        """
        pdf_links = {}

        try:
            page_token = None
            while True:
                results = self.drive_service.service.files().list(
                    q=f"'{drive_folder_id}' in parents and mimeType='application/pdf' and trashed=false",
                    fields="nextPageToken, files(id, name, webViewLink)",
                    pageSize=1000,
                    pageToken=page_token,
                    includeItemsFromAllDrives=True,
                    supportsAllDrives=True
                ).execute()

                files = results.get('files', [])
                for file in files:
                    filename = file['name']

                    # Extract page range from filename
                    # Expected format: "3-4.pdf", "12-18.pdf", etc.
                    if filename.endswith('.pdf'):
                        page_range = filename[:-4]  # Remove .pdf extension

                        # Validate it's a page range format (digits-digits)
                        if re.match(r'^\d+-\d+$', page_range):
                            pdf_links[page_range] = file['webViewLink']

                page_token = results.get('nextPageToken')
                if not page_token:
                    break

            return pdf_links

        except Exception as e:
            raise Exception(f"Failed to fetch PDFs from Drive folder: {str(e)}")

    def get_word_file_from_drive(self, drive_file_id: str, download_path: str) -> str:
        """
        Download a Word document from Google Drive

        Args:
            drive_file_id: Google Drive file ID of the Word document
            download_path: Local path to save the downloaded file

        Returns:
            Path to downloaded file
        """
        try:
            # Get file metadata
            file_metadata = self.drive_service.service.files().get(
                fileId=drive_file_id,
                fields='name, mimeType',
                supportsAllDrives=True
            ).execute()

            # Ensure it's a Word document
            mime_type = file_metadata.get('mimeType', '')
            if 'word' not in mime_type.lower() and not file_metadata['name'].endswith('.docx'):
                raise Exception(f"File is not a Word document: {file_metadata['name']}")

            # Download file
            request = self.drive_service.service.files().get_media(fileId=drive_file_id)

            import io
            fh = io.BytesIO()
            from googleapiclient.http import MediaIoBaseDownload
            downloader = MediaIoBaseDownload(fh, request)

            done = False
            while not done:
                status, done = downloader.next_chunk()

            # Save to file
            Path(download_path).parent.mkdir(parents=True, exist_ok=True)
            with open(download_path, 'wb') as f:
                f.write(fh.getvalue())

            return download_path

        except Exception as e:
            raise Exception(f"Failed to download Word file from Drive: {str(e)}")

    def process_word_document(
        self,
        input_docx_path: str,
        pdf_links: Dict[str, str],
        output_docx_path: Optional[str] = None
    ) -> Dict:
        """
        Process Word document and attach Drive links to statements

        Args:
            input_docx_path: Path to input Word document
            pdf_links: Dict mapping page_range -> drive_webview_link
            output_docx_path: Path to save output (if None, overwrites input)

        Returns:
            Dict with processing statistics
        """
        if output_docx_path is None:
            output_docx_path = input_docx_path

        doc = Document(input_docx_path)

        stats = {
            'total_statements': 0,
            'linked_statements': 0,
            'unlinked_statements': 0,
            'statements': []
        }

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            # Try to parse statement with page number
            result = self.parse_statement_with_page_number(text)

            if result:
                stats['total_statements'] += 1
                page_range = result['page_range']
                clean_text = result['clean_text']
                format_type = result.get('format', 'unknown')

                # Check if we have a Drive link for this page range
                if page_range in pdf_links:
                    drive_link = pdf_links[page_range]
                    stats['linked_statements'] += 1

                    # Clear existing text
                    paragraph.clear()

                    if format_type == 'doctor_statement':
                        # DOCTOR FORMAT: "Doctor Name - File_ID" becomes a hyperlink
                        # The entire clean_text is the hyperlink
                        add_hyperlink(
                            paragraph,
                            drive_link,
                            clean_text,
                            color='0563C1',  # Standard hyperlink blue
                            underline=True,
                            bold=True
                        )

                    elif format_type == 'flexible_page_range':
                        # FLEXIBLE FORMAT: Clean text becomes hyperlink (page range removed)
                        # Works with ANY format: "Page No 3-4", "(3-4)", "Page_No_3-4", etc.
                        add_hyperlink(
                            paragraph,
                            drive_link,
                            clean_text,
                            color='0563C1',  # Standard hyperlink blue
                            underline=True,
                            bold=True
                        )

                    elif format_type == 'page_no_statement':
                        # LEGACY FORMAT: Keep the header + rest of text structure
                        full_text_hash = result.get('parts', {})

                        # 1. Add the "clean text" (Header) as the hyperlink
                        add_hyperlink(
                            paragraph,
                            drive_link,
                            clean_text,
                            color='0563C1',  # Standard hyperlink blue
                            underline=True,
                            bold=True
                        )

                        # 2. Append the rest of the paragraph (Page No + Body) as regular text
                        rest_of_text = full_text_hash.get('match_text', '') + full_text_hash.get('post_text', '')
                        if rest_of_text:
                            paragraph.add_run(" " + rest_of_text)

                    stats['statements'].append({
                        'page_range': page_range,
                        'text': clean_text,
                        'status': 'linked',
                        'drive_link': drive_link
                    })
                else:
                    stats['unlinked_statements'] += 1

                    stats['statements'].append({
                        'page_range': page_range,
                        'text': clean_text,
                        'status': 'not_found',
                        'reason': f'No PDF found for pages {page_range}'
                    })

        # Save document
        doc.save(output_docx_path)

        return stats

    def upload_processed_word_to_drive(
        self,
        local_docx_path: str,
        drive_folder_id: str,
        filename: Optional[str] = None
    ) -> Tuple[str, str]:
        """
        Upload processed Word document back to Google Drive

        Args:
            local_docx_path: Path to local Word document
            drive_folder_id: Google Drive folder ID to upload to
            filename: Optional filename (uses local filename if None)

        Returns:
            Tuple of (file_id, webview_link)
        """
        if filename is None:
            filename = Path(local_docx_path).name

        return self.drive_service.upload_file(
            local_docx_path,
            drive_folder_id,
            filename
        )

    def process_from_drive_folder(
        self,
        drive_folder_id: str,
        word_file_id: str,
        output_folder_id: Optional[str] = None,
        temp_dir: str = '/tmp'
    ) -> Dict:
        """
        Complete workflow: Process Word document with PDFs from same Drive folder

        Args:
            drive_folder_id: Google Drive folder ID containing PDFs
            word_file_id: Google Drive file ID of Word document
            output_folder_id: Optional Drive folder to upload result (uses input folder if None)
            temp_dir: Temporary directory for processing

        Returns:
            Dict with processing results including file_id and webview_link of output
        """
        if output_folder_id is None:
            output_folder_id = drive_folder_id

        # Create temp paths
        temp_input = Path(temp_dir) / f"input_{word_file_id}.docx"
        temp_output = Path(temp_dir) / f"output_{word_file_id}.docx"

        try:
            # Step 1: Get PDFs from Drive folder
            pdf_links = self.get_pdfs_from_drive_folder(drive_folder_id)

            if not pdf_links:
                raise Exception("No PDF files found in Drive folder")

            # Step 2: Download Word document
            self.get_word_file_from_drive(word_file_id, str(temp_input))

            # Step 3: Process document
            stats = self.process_word_document(
                str(temp_input),
                pdf_links,
                str(temp_output)
            )

            # Step 4: Upload processed document
            output_file_id, output_webview_link = self.upload_processed_word_to_drive(
                str(temp_output),
                output_folder_id,
                filename=Path(temp_input).name  # Keep original filename
            )

            # Add upload info to stats
            stats['output_file_id'] = output_file_id
            stats['output_webview_link'] = output_webview_link
            stats['pdf_count'] = len(pdf_links)

            return stats

        finally:
            # Cleanup temp files
            for path in [temp_input, temp_output]:
                if path.exists():
                    path.unlink()

    def process_local_workflow(
        self,
        word_file_path: str,
        pdf_folder_path: str,
        output_word_path: Optional[str] = None,
        upload_to_drive: bool = False,
        drive_folder_id: Optional[str] = None
    ) -> Dict:
        """
        Process Word document with local PDFs and optionally upload to Drive

        Args:
            word_file_path: Path to local Word document
            pdf_folder_path: Path to folder containing PDF files (named as page ranges)
            output_word_path: Path to save output (if None, overwrites input)
            upload_to_drive: Whether to upload PDFs and processed doc to Drive
            drive_folder_id: Google Drive folder ID (required if upload_to_drive=True)

        Returns:
            Dict with processing statistics
        """
        pdf_links = {}

        if upload_to_drive:
            if not drive_folder_id:
                raise Exception("drive_folder_id required when upload_to_drive=True")

            # Upload PDFs to Drive first
            pdf_folder = Path(pdf_folder_path)
            pdf_files = sorted(pdf_folder.glob('*-*.pdf'))  # Files with page range format

            for pdf_file in pdf_files:
                page_range = pdf_file.stem  # e.g., "3-4"

                # Upload to Drive
                file_id, webview_link = self.drive_service.upload_file(
                    str(pdf_file),
                    drive_folder_id,
                    pdf_file.name
                )

                pdf_links[page_range] = webview_link
        else:
            # For local processing, we'd need local file links
            # This is not typical for production, but included for completeness
            raise NotImplementedError(
                "Local-only processing requires uploading PDFs to Drive. "
                "Set upload_to_drive=True and provide drive_folder_id."
            )

        # Process document
        stats = self.process_word_document(
            word_file_path,
            pdf_links,
            output_word_path
        )

        # Upload processed Word document if requested
        if upload_to_drive and output_word_path:
            output_file_id, output_webview_link = self.upload_processed_word_to_drive(
                output_word_path,
                drive_folder_id
            )

            stats['output_file_id'] = output_file_id
            stats['output_webview_link'] = output_webview_link

        stats['pdf_count'] = len(pdf_links)

        return stats
