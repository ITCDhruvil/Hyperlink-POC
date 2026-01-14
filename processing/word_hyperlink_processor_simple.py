"""
Simple Word Hyperlink Processor

SIMPLE PATTERN:
- Statements have page numbers at the END
- Example: "04/05/13. Progress Note. US HEALTHWORKS. 4-9"
- Extract page number, make whole statement a hyperlink
- Result: "04/05/13. Progress Note. US HEALTHWORKS." (linked to 4-9.pdf)
"""
import re
import copy
from typing import Optional, Dict
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from .drive_utils import get_drive_service


def _normalize_page_spec(s: str) -> str:
    s = (s or '').replace('\u00a0', ' ').strip()
    s = s.replace('–', '-').replace('—', '-').replace('‑', '-')
    # Normalize spaces around dashes in page ranges: "1 - 2" -> "1-2"
    s = re.sub(r'(\d+)\s*-\s*(\d+)', r'\1-\2', s)
    s = re.sub(r'\s+', ' ', s)
    s = re.sub(r'\s*,\s*', ', ', s)
    s = s.strip()
    if not _is_page_spec(s):
        return s

    parts = [p.strip() for p in s.split(',') if p.strip()]
    normalized_parts = []
    for p in parts:
        m = re.fullmatch(r'(\d+)(?:\s*-\s*(\d+))?', p)
        if not m:
            normalized_parts.append(p)
            continue
        a = str(int(m.group(1)))
        if m.group(2):
            b = str(int(m.group(2)))
            normalized_parts.append(f"{a}-{b}")
        else:
            normalized_parts.append(a)

    return ', '.join(normalized_parts)


def _is_page_spec(s: str) -> bool:
    return bool(re.fullmatch(r'\d+(?:\s*-\s*\d+)?(?:\s*,\s*\d+(?:\s*-\s*\d+)?)*', (s or '').strip()))


def add_hyperlink(paragraph, url, text, color='0563C1', underline=True, bold=False, font_name: str = 'Times New Roman', font_size_pt: int = 12):
    """Add a hyperlink to a paragraph"""
    part = paragraph.part
    r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Hyperlink style
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)
    rPr.append(rFonts)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(int(font_size_pt) * 2))
    rPr.append(sz)

    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(int(font_size_pt) * 2))
    rPr.append(szCs)

    # Color
    c = OxmlElement('w:color')
    c.set(qn('w:val'), color)
    rPr.append(c)

    # Underline
    if underline:
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)

    new_run.append(rPr)

    # Text
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)

    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


class WordHyperlinkProcessorSimple:
    """
    Simple processor for statements with page numbers at the end
    """

    def __init__(self):
        self.drive_service = None

    def _get_drive_service(self):
        if self.drive_service is None:
            self.drive_service = get_drive_service()
        return self.drive_service

    def extract_page_ranges_from_document(self, doc: Document) -> list:
        """
        Extract all page ranges from a Word document.

        Scans the document for statements (lines starting with dates) and
        extracts their page numbers. Returns list of page ranges in order.

        Args:
            doc: python-docx Document object

        Returns:
            List of page ranges as strings (e.g., ['1-2', '4', '7-8', ...])

        Example:
            processor = WordHyperlinkProcessorSimple()
            doc = Document('patient_record.docx')
            ranges = processor.extract_page_ranges_from_document(doc)
            # Returns: ['1-2', '4', '7-8', '11-15', ...]

            # Format for PDF split:
            formatted = ';'.join(ranges)
            # Returns: '1-2;4;7-8;11-15'
        """
        page_ranges = []
        seen_ranges = set()

        # Get all paragraphs (body + tables)
        all_paragraphs = []
        for p in doc.paragraphs:
            all_paragraphs.append(p.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        all_paragraphs.append(p.text)

        # Process each paragraph
        for text in all_paragraphs:
            if not text or not text.strip():
                continue

            # Check if it's a statement line
            if not self.is_statement_line(text):
                continue

            # Extract page range
            result = self.parse_statement_with_page_number(text)
            if not result or not result.get('page_range'):
                continue

            page_range = result['page_range']

            # Avoid duplicates
            if page_range not in seen_ranges:
                page_ranges.append(page_range)
                seen_ranges.add(page_range)

        return page_ranges

    def extract_page_ranges_from_file(self, docx_path: str) -> list:
        """
        Extract page ranges from a Word document file path.

        Convenience method that opens the document and extracts ranges.

        Args:
            docx_path: Path to .docx file

        Returns:
            List of page ranges as strings

        Example:
            processor = WordHyperlinkProcessorSimple()
            ranges = processor.extract_page_ranges_from_file('patient.docx')
            formatted = ';'.join(ranges)  # For PDF split
        """
        doc = Document(docx_path)
        return self.extract_page_ranges_from_document(doc)

    def extract_patient_name_from_document(self, doc: Document) -> Optional[str]:
        """
        Extract patient name from document header
        Looks for: "PATIENT NAME: CARL MAYFIELD"
        """
        for paragraph in doc.paragraphs[:10]:
            text = paragraph.text.strip()

            # Pattern: "PATIENT NAME: CARL MAYFIELD"
            match = re.search(r'PATIENT\s+NAME:?\s+(.+)', text, re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                # Convert to folder format
                name = name.title().replace(' ', '_')
                return name

            # Pattern: "Name: Carl Mayfield"
            match = re.search(r'^Name:?\\s+(.+)', text, re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                name = name.title().replace(' ', '_')
                return name

        return None

    def parse_statement_with_page_number(self, text: str) -> Optional[Dict]:
        """
        Parse statement to extract page number

        Pattern: "DATE.  DESCRIPTION. PAGE_NUMBERS ADDITIONAL_TEXT"
        Example: "06/19/25.  From EMANATE HEALTH...  Attestation. 1-2 Attesting to 422 pages..."

        The page numbers appear after the last period in the description, followed by optional text.

        Returns:
            {
                'page_range': '1-2',
                'header_text': '06/19/25.  From EMANATE HEALTH...  Attestation.',
                'remainder_text': 'Attesting to 422 pages...',
                'original_text': 'original...'
            }
        """
        original_text = (text or '').replace('\u00a0', ' ').replace('–', '-').replace('—', '-').replace('‑', '-')

        # Pattern: digit or range (1-2) or multiple ranges (25-29, 31-35, 40)
        # We select the candidate closest to the end of the statement and require it to be
        # followed by either end-of-line / trailing punctuation OR a dot that starts the next
        # sentence/statement (often ". <letters>").
        page_spec_pattern = re.compile(
            r'(?P<pages>\d+(?:\s*-\s*\d+)?(?:\s*,\s*\d+(?:\s*-\s*\d+)?)*)'
        )

        best = None
        best_rank = None  # tuple: (-complexity, distance_to_end)

        for m in page_spec_pattern.finditer(original_text):
            start, end = m.start('pages'), m.end('pages')

            # Reject if part of a date like "06/19/25"
            if end < len(original_text) and original_text[end] == '/':
                continue
            if start > 0 and original_text[start - 1] == '/':
                continue

            # Reject if part of a decimal number like "0.9" or "164.5"
            if start > 0 and original_text[start - 1] == '.':
                continue

            # Reject if part of numbered list like "3) Renal mass"
            if end < len(original_text) and original_text[end] == ')':
                continue

            pages_raw = (m.group('pages') or '').strip()
            if not _is_page_spec(pages_raw):
                continue

            # Check if this appears after a period or colon (common separators)
            if start > 0:
                before = original_text[:start].rstrip(' ')
                if not before:
                    continue
                if not (before.endswith('.') or before.endswith(':') or before.endswith(' ')):
                    continue

                context_before = before[-15:] if len(before) >= 15 else before
                if re.search(r'(SpO2|BP|Pulse|Temp|Weight|RR|P|T):\s*$', context_before, re.IGNORECASE):
                    continue

                # Reject common trailing numeric fields (these are not page specs)
                # Example: "Sessions: 22." or "Visits: 12."
                context_before_long = before[-30:] if len(before) >= 30 else before
                if re.search(r'(Sessions?|Visits?|Appts?|Appointments?|Units?):\s*$', context_before_long, re.IGNORECASE):
                    continue

            # Boundary check: pages should run until end, trailing punctuation, or ". <letters>"
            j = end
            while j < len(original_text) and original_text[j].isspace():
                j += 1

            boundary_ok = False
            if j >= len(original_text):
                boundary_ok = True
            else:
                ch = original_text[j]
                # allow trailing punctuation at end of the statement
                if ch in '.;,:' and j == len(original_text) - 1:
                    boundary_ok = True
                elif ch in '.;,:' and j + 1 < len(original_text):
                    # If this is ". <letters>" treat it as boundary (next statement/sentence starts)
                    k = j + 1
                    while k < len(original_text) and original_text[k].isspace():
                        k += 1
                    if k < len(original_text) and original_text[k].isalpha():
                        boundary_ok = True

            if not boundary_ok:
                continue

            distance_to_end = len(original_text) - end
            segments = pages_raw.count(',') + 1
            has_range = 1 if '-' in pages_raw else 0
            complexity = (segments * 10) + has_range + len(pages_raw)
            rank = (-complexity, distance_to_end)
            if best is None or (best_rank is not None and rank < best_rank):
                best = m
                best_rank = rank

        if best:
            page_spec_raw = (best.group('pages') or '').strip()
            page_spec = _normalize_page_spec(page_spec_raw)

            header_text = original_text[:best.start('pages')]
            remainder_text = original_text[best.end('pages'):]

            return {
                'page_range': page_spec,
                'header_text': header_text,
                'remainder_text': remainder_text,
                'original_text': original_text,
                'pages_span': (best.start('pages'), best.end('pages')),
            }

        return None

    def is_statement_line(self, text: str) -> bool:
        """
        Check if line is a statement (starts with date)

        Patterns:
        - "06/19/25.  From EMANATE HEALTH..."
        - "08/26/22.  EMANATE HEALTH..."

        Must start with DATE (no numbering prefix required)
        """
        text = (text or '').replace('\u00a0', ' ').strip()
        text = text.replace('–', '-').replace('—', '-').replace('‑', '-')
        # Pattern: starts with either:
        # - MM/DD/YY. (or MM/DD/YYYY.)
        # - MM/DD/YY-MM/DD/YY. (or with spaces around '-')
        return bool(
            re.match(
                r'^\d{1,2}/\d{1,2}/\d{2,4}(?:\s*-\s*\d{1,2}/\d{1,2}/\d{2,4})?\.\s+',
                text,
            )
        )

    def _split_run(self, run, split_at: int):
        if split_at <= 0 or split_at >= len(run.text or ''):
            return None
        tail_text = (run.text or '')[split_at:]
        run.text = (run.text or '')[:split_at]
        new_r = OxmlElement('w:r')
        if run._r.rPr is not None:
            new_r.append(copy.deepcopy(run._r.rPr))
        t = OxmlElement('w:t')
        if tail_text.startswith(' ') or tail_text.endswith(' '):
            t.set(qn('xml:space'), 'preserve')
        t.text = tail_text
        new_r.append(t)
        run._r.addnext(new_r)
        return run._parent.add_run()._r.getprevious()

    def _link_statement_in_paragraph(self, paragraph, pages_span: tuple, url: str) -> None:
        """Option B: remove the page-range text and hyperlink+bold the header text.

        This implementation edits the paragraph XML in-place to preserve spacing and existing
        formatting for content we do not touch.
        """
        start_idx, end_idx = pages_span
        if start_idx < 0 or end_idx <= start_idx:
            return

        p = paragraph._p

        def _iter_run_elements():
            for el in list(p):
                if el.tag.endswith('}r'):
                    yield el

        def _run_text(run_el) -> str:
            texts = []
            for t in run_el.findall(qn('w:t')):
                texts.append(t.text or '')
            return ''.join(texts)

        def _set_run_text(run_el, new_text: str) -> None:
            ts = run_el.findall(qn('w:t'))
            if not ts:
                t = OxmlElement('w:t')
                run_el.append(t)
                ts = [t]
            # keep first, remove the rest
            for extra in ts[1:]:
                run_el.remove(extra)
            t0 = ts[0]
            if new_text.startswith(' ') or new_text.endswith(' ') or '  ' in new_text:
                t0.set(qn('xml:space'), 'preserve')
            t0.text = new_text

        def _split_run_el(run_el, offset: int):
            txt = _run_text(run_el)
            if offset <= 0 or offset >= len(txt):
                return run_el
            left = txt[:offset]
            right = txt[offset:]
            _set_run_text(run_el, left)
            new_el = copy.deepcopy(run_el)
            _set_run_text(new_el, right)
            p.insert(p.index(run_el) + 1, new_el)
            return new_el

        # Build run spans
        runs = []
        acc = 0
        for r in _iter_run_elements():
            txt = _run_text(r)
            if not txt:
                continue
            runs.append({'el': r, 'start': acc, 'end': acc + len(txt)})
            acc += len(txt)

        if not runs:
            return

        # Split at boundaries so deletion is whole-run
        for item in list(runs):
            r = item['el']
            s = item['start']
            e = item['end']
            if s < start_idx < e:
                _split_run_el(r, start_idx - s)
                break

        # Recompute spans after potential split
        runs = []
        acc = 0
        for r in _iter_run_elements():
            txt = _run_text(r)
            if not txt:
                continue
            runs.append({'el': r, 'start': acc, 'end': acc + len(txt)})
            acc += len(txt)

        for item in list(runs):
            r = item['el']
            s = item['start']
            e = item['end']
            if s < end_idx < e:
                _split_run_el(r, end_idx - s)
                break

        # Recompute spans again
        runs = []
        acc = 0
        for r in _iter_run_elements():
            txt = _run_text(r)
            if not txt:
                continue
            runs.append({'el': r, 'start': acc, 'end': acc + len(txt)})
            acc += len(txt)

        # Remove page-range runs
        for item in list(runs):
            if item['start'] >= start_idx and item['end'] <= end_idx:
                p.remove(item['el'])

        # Collect header runs (everything before start_idx)
        header_runs = []
        acc = 0
        for r in _iter_run_elements():
            txt = _run_text(r)
            if not txt:
                continue
            if acc + len(txt) > start_idx:
                if acc < start_idx:
                    _split_run_el(r, start_idx - acc)
                    # After split, the current run is header part; re-iterate fresh below
                break
            header_runs.append(r)
            acc += len(txt)

        # refresh header_runs after possible split
        header_runs = []
        acc = 0
        for r in _iter_run_elements():
            txt = _run_text(r)
            if not txt:
                continue
            if acc >= start_idx:
                break
            if acc + len(txt) > start_idx:
                break
            header_runs.append(r)
            acc += len(txt)

        if not header_runs:
            return

        # Create hyperlink and move header runs into it
        r_id = paragraph.part.relate_to(
            url,
            'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',
            is_external=True,
        )
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        first_header_run = header_runs[0]
        insert_pos = p.index(first_header_run)

        def _ensure_rpr(run_el):
            rpr = run_el.find(qn('w:rPr'))
            if rpr is None:
                rpr = OxmlElement('w:rPr')
                run_el.insert(0, rpr)
            return rpr

        def _set_bold_and_color(run_el):
            rpr = _ensure_rpr(run_el)
            if rpr.find(qn('w:b')) is None:
                b = OxmlElement('w:b')
                rpr.append(b)

            u = rpr.find(qn('w:u'))
            if u is None:
                u = OxmlElement('w:u')
                rpr.append(u)
            u.set(qn('w:val'), 'single')

            color = rpr.find(qn('w:color'))
            if color is None:
                color = OxmlElement('w:color')
                rpr.append(color)
            color.set(qn('w:val'), '0000FF')

        for r in header_runs:
            _set_bold_and_color(r)
            hyperlink.append(r)

        p.insert(insert_pos, hyperlink)

    def _iter_all_paragraphs(self, doc: Document):
        # Yield normal body paragraphs
        for p in doc.paragraphs:
            yield p

        # Yield paragraphs inside tables (common in generated Word docs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p

    def _apply_default_font(self, doc: Document, font_name: str = 'Times New Roman', font_size_pt: int = 12) -> None:
        normal = doc.styles['Normal']
        normal.font.name = font_name
        normal.font.size = Pt(font_size_pt)
        if normal._element.rPr is not None and normal._element.rPr.rFonts is not None:
            normal._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        for p in self._iter_all_paragraphs(doc):
            for run in p.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size_pt)
                if run._element is not None and run._element.rPr is not None and run._element.rPr.rFonts is not None:
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    def _remove_paragraph_borders(self, paragraph) -> None:
        pPr = paragraph._p.pPr
        if pPr is None:
            return
        pBdr = pPr.find(qn('w:pBdr'))
        if pBdr is not None:
            pPr.remove(pBdr)

    def _normalize_tables(self, doc: Document) -> None:
        for table in doc.tables:
            try:
                table.style = 'Table Grid'
            except Exception:
                pass

            tbl = table._tbl
            tblPr = tbl.tblPr
            if tblPr is None:
                continue

            existing_borders = tblPr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tblPr.remove(existing_borders)

            borders = OxmlElement('w:tblBorders')

            def _add_border(tag: str):
                el = OxmlElement(tag)
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), '4')
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), 'auto')
                return el

            borders.append(_add_border('w:top'))
            borders.append(_add_border('w:left'))
            borders.append(_add_border('w:bottom'))
            borders.append(_add_border('w:right'))
            borders.append(_add_border('w:insideH'))
            borders.append(_add_border('w:insideV'))

            tblPr.append(borders)

            # Clear cell shading/borders that can appear as thick black/gray bands
            for row in table.rows:
                for cell in row.cells:
                    tcPr = cell._tc.tcPr
                    if tcPr is None:
                        continue

                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        tcPr.remove(shd)

                    tcBorders = tcPr.find(qn('w:tcBorders'))
                    if tcBorders is not None:
                        tcPr.remove(tcBorders)

    def get_pdfs_from_drive_folder(self, drive_folder_id: str) -> Dict[str, str]:
        """Get all PDF files from Drive folder"""
        pdf_links = {}

        try:
            drive_service = self._get_drive_service()
            page_token = None
            while True:
                results = drive_service.service.files().list(
                    q=f"'{drive_folder_id}' in parents and mimeType='application/pdf' and trashed=false",
                    fields="nextPageToken, files(id, name, webViewLink)",
                    pageSize=1000,
                    pageToken=page_token,
                    includeItemsFromAllDrives=True,
                    supportsAllDrives=True
                ).execute()

                files = results.get('files', [])
                for file in files:
                    filename = file['name']
                    if not filename.lower().endswith('.pdf'):
                        continue

                    stem = filename[:-4]
                    key = _normalize_page_spec(stem)
                    if _is_page_spec(key):
                        pdf_links[key] = file['webViewLink']

                page_token = results.get('nextPageToken')
                if not page_token:
                    break

            return pdf_links

        except Exception as e:
            raise Exception(f"Failed to fetch PDFs: {str(e)}")

    def process_word_document(
        self,
        input_docx_path: str,
        pdf_links: Dict[str, str],
        output_docx_path: Optional[str] = None
    ) -> Dict:
        """
        Process Word document - SIMPLE PATTERN

        Logic:
        1. Find statement with page number at end
        2. Extract page number
        3. Replace whole statement with hyperlink (without page number)
        """
        if output_docx_path is None:
            output_docx_path = input_docx_path

        doc = Document(input_docx_path)

        stats = {
            'total_statements': 0,
            'linked_statements': 0,
            'unlinked_statements': 0,
            'statements': []
        }

        for paragraph in self._iter_all_paragraphs(doc):
            raw_text = paragraph.text
            if not raw_text or not raw_text.strip():
                continue

            normalized_text = (raw_text or '').replace('\u00a0', ' ').replace('–', '-').replace('—', '-').replace('‑', '-')
            if not self.is_statement_line(normalized_text):
                continue

            statement_info = self.parse_statement_with_page_number(normalized_text)
            if not statement_info:
                continue

            page_range = statement_info['page_range']
            pages_span = statement_info.get('pages_span')
            if not pages_span:
                continue

            stats['total_statements'] += 1

            drive_link = pdf_links.get(page_range)
            if not drive_link:
                stats['unlinked_statements'] += 1
                stats['statements'].append({
                    'page_range': page_range,
                    'text': statement_info.get('header_text', ''),
                    'status': 'not_found',
                    'reason': f'No PDF found for pages {page_range}'
                })
                continue

            try:
                self._remove_paragraph_borders(paragraph)
                self._link_statement_in_paragraph(paragraph, pages_span, drive_link)
            except Exception:
                stats['unlinked_statements'] += 1
                stats['statements'].append({
                    'page_range': page_range,
                    'text': statement_info.get('header_text', ''),
                    'status': 'not_found',
                    'reason': f'Failed to apply hyperlink for pages {page_range}'
                })
                continue

            stats['linked_statements'] += 1
            stats['statements'].append({
                'page_range': page_range,
                'text': statement_info.get('header_text', ''),
                'status': 'linked',
                'drive_link': drive_link
            })

        self._apply_default_font(doc, font_name='Times New Roman', font_size_pt=12)
        self._normalize_tables(doc)
        doc.save(output_docx_path)
        return stats
